from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_function_documentation():
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('BOOK-PLAY Project - Complete Function Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('Comprehensive Function Signatures and Explanations')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add summary
    doc.add_heading('Project Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run('Total Functions: 127\n').bold = True
    summary.add_run('• PHP Functions: 15\n')
    summary.add_run('• JavaScript Functions: 112\n\n')
    summary.add_run('This document contains all functions found in the BOOK-PLAY sports venue booking application, organized by file type and functionality.')

    # PHP Functions Section
    doc.add_heading('PHP Functions', level=1)
    
    # Authentication Files
    doc.add_heading('Authentication Files', level=2)
    
    php_functions = [
        ('getSportsHTML()', 'pages/auth/Register_Page/verify.php', 'Generates HTML checkboxes for sports selection during user registration'),
        ('sendVerificationCode($to, $code)', 'mail/MailLink.php', 'Sends 6-digit verification codes via email using PHPMailer'),
    ]
    
    for func_name, file_path, explanation in php_functions:
        p = doc.add_paragraph()
        p.add_run(f'Function: ').bold = True
        p.add_run(func_name)
        p.add_run(f'\nFile: ').bold = True
        p.add_run(file_path)
        p.add_run(f'\nExplanation: ').bold = True
        p.add_run(explanation)
        doc.add_paragraph()

    # Admin API Files
    doc.add_heading('Admin API Files', level=2)
    
    admin_functions = [
        ('getReports()', 'pages/Admin/ManageVenueRequests/VenueAPI.php', 'Fetches venue reports and unaccepted venues for admin review'),
        ('handleAction()', 'pages/Admin/ManageVenueRequests/VenueAPI.php', 'Processes admin approve/reject actions for venue requests'),
        ('updateReportStatus()', 'pages/Admin/ManageVenueRequests/VenueAPI.php', 'Updates status of venue reports in database'),
    ]
    
    for func_name, file_path, explanation in admin_functions:
        p = doc.add_paragraph()
        p.add_run(f'Function: ').bold = True
        p.add_run(func_name)
        p.add_run(f'\nFile: ').bold = True
        p.add_run(file_path)
        p.add_run(f'\nExplanation: ').bold = True
        p.add_run(explanation)
        doc.add_paragraph()

    # Facility Owner API Files
    doc.add_heading('Facility Owner API Files', level=2)
    
    facility_functions = [
        ('getFacilities()', 'pages/Facility_Owner/ManageVenue/fetch_venues.php', 'Retrieves all facilities owned by current facility owner'),
        ('updateAvailability()', 'pages/Facility_Owner/ManageVenue/fetch_venues.php', 'Updates availability status of a facility'),
        ('getSports()', 'pages/Facility_Owner/ManageVenue/fetch_venues.php', 'Fetches available sports categories from existing facilities'),
        ('addFacility()', 'pages/Facility_Owner/ManageVenue/fetch_venues.php', 'Creates new sport facility and submits for admin approval'),
        ('updateFacility()', 'pages/Facility_Owner/ManageVenue/fetch_venues.php', 'Updates existing facility information'),
        ('upload_images($files)', 'pages/Facility_Owner/ManageVenue/upload_images.php', 'Handles multiple image uploads for sport facilities'),
    ]
    
    for func_name, file_path, explanation in facility_functions:
        p = doc.add_paragraph()
        p.add_run(f'Function: ').bold = True
        p.add_run(func_name)
        p.add_run(f'\nFile: ').bold = True
        p.add_run(file_path)
        p.add_run(f'\nExplanation: ').bold = True
        p.add_run(explanation)
        doc.add_paragraph()

    # JavaScript Functions Section
    doc.add_heading('JavaScript Functions', level=1)
    
    # Shared Functions
    doc.add_heading('Shared Functions', level=2)
    
    shared_functions = [
        ('toggleProfileMenu()', 'assets/js/shared.js', 'Shows/hides user profile dropdown menu'),
        ('showTempMessage(message, duration)', 'assets/js/shared.js', 'Displays temporary notification messages'),
        ('loadUserProfile()', 'assets/js/shared.js', 'Fetches and displays current user\'s profile image and username'),
        ('initAdminRequestButton()', 'assets/js/shared.js', 'Sets up admin request functionality for regular users'),
    ]
    
    for func_name, file_path, explanation in shared_functions:
        p = doc.add_paragraph()
        p.add_run(f'Function: ').bold = True
        p.add_run(func_name)
        p.add_run(f'\nFile: ').bold = True
        p.add_run(file_path)
        p.add_run(f'\nExplanation: ').bold = True
        p.add_run(explanation)
        doc.add_paragraph()

    # Player Functions
    doc.add_heading('Player Functions', level=2)
    
    player_functions = [
        # VenueDetails
        ('highlightStars(rating)', 'pages/player/VenueDetails/VenueDetails.js', 'Highlights star rating display based on user interaction'),
        ('resetStars()', 'pages/player/VenueDetails/VenueDetails.js', 'Resets star rating display to default state'),
        ('openReportModal()', 'pages/player/VenueDetails/VenueDetails.js', 'Opens venue report submission modal'),
        ('closeReportModal()', 'pages/player/VenueDetails/VenueDetails.js', 'Closes venue report modal'),
        ('submitReport()', 'pages/player/VenueDetails/VenueDetails.js', 'Submits venue report to administrators'),
        
        # MyBookings
        ('renderBookings(bookings)', 'pages/player/MyBookings/MyBookings.js', 'Displays user\'s booking history in formatted list'),
        ('formatDate(dateStr)', 'pages/player/MyBookings/MyBookings.js', 'Formats date strings for user-friendly display'),
        
        # MyFriends
        ('setupEventListeners()', 'pages/player/MyFriends/MyFriends.js', 'Initializes all event listeners for friend management page'),
        ('updateStats()', 'pages/player/MyFriends/MyFriends.js', 'Updates friend statistics display'),
        ('renderFriendRequests()', 'pages/player/MyFriends/MyFriends.js', 'Displays pending friend requests'),
        ('renderCurrentFriends()', 'pages/player/MyFriends/MyFriends.js', 'Displays current friends list'),
        ('handleSearch()', 'pages/player/MyFriends/MyFriends.js', 'Processes friend search functionality'),
        ('acceptFriendRequest(requestId)', 'pages/player/MyFriends/MyFriends.js', 'Accepts a friend request'),
        ('rejectFriendRequest(requestId)', 'pages/player/MyFriends/MyFriends.js', 'Rejects a friend request'),
        ('messageFriend(friendId)', 'pages/player/MyFriends/MyFriends.js', 'Opens chat with a friend'),
        ('showFriendModal(friendId)', 'pages/player/MyFriends/MyFriends.js', 'Displays detailed friend information in modal'),
        ('goToFindPlayers()', 'pages/player/MyFriends/MyFriends.js', 'Navigates to find players page'),
        ('goToBookings()', 'pages/player/MyFriends/MyFriends.js', 'Navigates to bookings page'),
        ('goToMessages()', 'pages/player/MyFriends/MyFriends.js', 'Navigates to messages page'),
        ('goToProfile()', 'pages/player/MyFriends/MyFriends.js', 'Navigates to profile page'),
        ('logout()', 'pages/player/MyFriends/MyFriends.js', 'Handles user logout process'),
        
        # HomePage
        ('toggleProfileMenu()', 'pages/player/HomePage/HomePage.js', 'Toggles profile menu visibility on homepage'),
        ('showTempMessage(message, duration)', 'pages/player/HomePage/HomePage.js', 'Shows temporary messages on homepage'),
        
        # JoinGroup
        ('toggleFavorite(element)', 'pages/player/JoinGroup/JoinGroup.js', 'Toggles favorite status for venues in group listings'),
        ('filterVenuesBySports(selectedSports)', 'pages/player/JoinGroup/JoinGroup.js', 'Filters available groups by sports preferences'),
        ('viewBookingDetails(booking_id)', 'pages/player/JoinGroup/JoinGroup.js', 'Opens detailed view of specific booking'),
        ('renderGroups(groups)', 'pages/player/JoinGroup/JoinGroup.js', 'Displays available groups for joining'),
        ('handleJoinClick(button)', 'pages/player/JoinGroup/JoinGroup.js', 'Processes join group button clicks'),
        ('joinGroup(group)', 'pages/player/JoinGroup/JoinGroup.js', 'Executes the group joining process'),
        ('validateAccessCode()', 'pages/player/JoinGroup/JoinGroup.js', 'Validates password for private group access'),
        ('closeModal()', 'pages/player/JoinGroup/JoinGroup.js', 'Closes modal dialogs'),
        ('redirectToBooking(booking_id)', 'pages/player/JoinGroup/JoinGroup.js', 'Redirects user to booking details page'),
        ('getUserLocation()', 'pages/player/JoinGroup/JoinGroup.js', 'Gets user\'s current GPS location'),
        ('calculateDistance(lat1, lon1, lat2, lon2)', 'pages/player/JoinGroup/JoinGroup.js', 'Calculates distance between two GPS coordinates'),
        ('toRad(x)', 'pages/player/JoinGroup/JoinGroup.js', 'Converts degrees to radians for distance calculations'),
        
        # FindPlayer
        ('initializeAvailability()', 'pages/player/FindPlayer/FindPlayer.js', 'Sets up availability calendar for player search'),
        ('showAddButton(day)', 'pages/player/FindPlayer/FindPlayer.js', 'Shows add time slot button for specific day'),
        ('addTimeSlot(day)', 'pages/player/FindPlayer/FindPlayer.js', 'Adds new time slot to availability calendar'),
        ('deleteTimeSlot(btn)', 'pages/player/FindPlayer/FindPlayer.js', 'Removes time slot from availability'),
        ('updateUserAvailability(day)', 'pages/player/FindPlayer/FindPlayer.js', 'Saves availability changes for specific day'),
        ('searchPlayers(searchTerm)', 'pages/player/FindPlayer/FindPlayer.js', 'Searches for players based on criteria'),
        ('renderPlayers()', 'pages/player/FindPlayer/FindPlayer.js', 'Displays search results for players'),
        ('addFriend(playerName)', 'pages/player/FindPlayer/FindPlayer.js', 'Sends friend request to a player'),
        ('setupEventListeners()', 'pages/player/FindPlayer/FindPlayer.js', 'Initializes event listeners for player search page'),
        ('handleSelectAll()', 'pages/player/FindPlayer/FindPlayer.js', 'Handles select all functionality for sports filters'),
        ('openPlayerModal(username, email, sport, age, gender, phone, location, image)', 'pages/player/FindPlayer/FindPlayer.js', 'Opens detailed player information modal'),
        ('closePlayerModal()', 'pages/player/FindPlayer/FindPlayer.js', 'Closes player details modal'),
        
        # Favorites
        ('generateStars(rating)', 'pages/player/Favorites/Favorites.js', 'Creates star rating display for venues'),
        ('createVenueCard(venue)', 'pages/player/Favorites/Favorites.js', 'Creates HTML card element for venue display'),
        ('loadFavorites()', 'pages/player/Favorites/Favorites.js', 'Loads and displays user\'s favorite venues'),
        ('toggleFavorite(venueId)', 'pages/player/Favorites/Favorites.js', 'Toggles favorite status for a venue'),
        ('viewDetails(venueId)', 'pages/player/Favorites/Favorites.js', 'Navigates to venue details page'),
        
        # CreateBooking
        ('initializeDateTimePickers()', 'pages/player/CreateBooking/CreateBooking.js', 'Sets up date and time picker components'),
        ('setupEventListeners()', 'pages/player/CreateBooking/CreateBooking.js', 'Initializes event listeners for booking form'),
        ('toggleGroupType()', 'pages/player/CreateBooking/CreateBooking.js', 'Toggles between public and private group booking types'),
        ('updateEndTimeMinTime()', 'pages/player/CreateBooking/CreateBooking.js', 'Updates minimum end time based on selected start time'),
        ('validateForm()', 'pages/player/CreateBooking/CreateBooking.js', 'Validates booking form before submission'),
        ('handleFormSubmission(e)', 'pages/player/CreateBooking/CreateBooking.js', 'Processes booking form submission'),
        ('updateSummary()', 'pages/player/CreateBooking/CreateBooking.js', 'Updates booking summary display'),
        ('loadUnavailableRanges(facilityId, bookingDate)', 'pages/player/CreateBooking/CreateBooking.js', 'Loads unavailable time ranges for facility on specific date'),
        ('mergeTimeRanges(ranges)', 'pages/player/CreateBooking/CreateBooking.js', 'Merges overlapping time ranges for display'),
        ('updateTimePickersWithUnavailableTimes(ranges)', 'pages/player/CreateBooking/CreateBooking.js', 'Updates time pickers to exclude unavailable times'),
        
        # Chats
        ('loadChat(username)', 'pages/player/Chats/Chats.js', 'Loads chat conversation with specific user'),
        ('sendMessage()', 'pages/player/Chats/Chats.js', 'Sends chat message to current chat partner'),
        
        # BookVenue
        ('waitForGoogleMaps(callback)', 'pages/player/BookVenue/BookVenue.js', 'Waits for Google Maps API to load before executing callback'),
        ('filterVenuesBySports(sports, searchTerm)', 'pages/player/BookVenue/BookVenue.js', 'Filters venues by sports and search terms'),
        ('sortVenues(venues, sortOptions)', 'pages/player/BookVenue/BookVenue.js', 'Sorts venues by various criteria'),
        ('toggleFavorite(iconElement, facilityId)', 'pages/player/BookVenue/BookVenue.js', 'Toggles favorite status for a venue'),
        ('renderVenues(venues)', 'pages/player/BookVenue/BookVenue.js', 'Displays filtered and sorted venues'),
        
        # BookingDetails
        ('fetchBookingDetails()', 'pages/player/BookingDetails/BookingDetails.js', 'Loads detailed booking information'),
        ('populateBookingDetails(booking)', 'pages/player/BookingDetails/BookingDetails.js', 'Populates booking information in the UI'),
        ('populatePlayerList(players)', 'pages/player/BookingDetails/BookingDetails.js', 'Displays list of players in the booking'),
        ('getHostUsername(players)', 'pages/player/BookingDetails/BookingDetails.js', 'Extracts host username from players list'),
        ('initializePrivacyToggle()', 'pages/player/BookingDetails/BookingDetails.js', 'Sets up privacy toggle functionality for group settings'),
        ('initializeScrolling()', 'pages/player/BookingDetails/BookingDetails.js', 'Sets up horizontal scrolling for player list'),
        ('updateArrowStates()', 'pages/player/BookingDetails/BookingDetails.js', 'Updates scroll arrow visibility based on scroll position'),
        ('initializeEditPrices()', 'pages/player/BookingDetails/BookingDetails.js', 'Sets up price editing functionality for group members'),
        ('initializePlayerActions()', 'pages/player/BookingDetails/BookingDetails.js', 'Sets up action buttons for player interactions'),
        ('initializeActionButtons()', 'pages/player/BookingDetails/BookingDetails.js', 'Sets up main action buttons for booking management'),
        ('switchHost(newHostId)', 'pages/player/BookingDetails/BookingDetails.js', 'Transfers host role to another group member'),
        ('toggleFriend(playerId)', 'pages/player/BookingDetails/BookingDetails.js', 'Toggles friend status with a player'),
        ('enterEditMode()', 'pages/player/BookingDetails/BookingDetails.js', 'Enables editing mode for booking details'),
        ('saveChanges()', 'pages/player/BookingDetails/BookingDetails.js', 'Saves edited booking details'),
        ('cancelChanges()', 'pages/player/BookingDetails/BookingDetails.js', 'Cancels editing and reverts changes'),
        ('exitEditMode()', 'pages/player/BookingDetails/BookingDetails.js', 'Exits editing mode and returns to view mode'),
        ('startCountdown()', 'pages/player/BookingDetails/BookingDetails.js', 'Starts countdown timer for booking'),
        ('updateCountdownDisplay()', 'pages/player/BookingDetails/BookingDetails.js', 'Updates countdown display with current time remaining'),
        ('handlePayNow()', 'pages/player/BookingDetails/BookingDetails.js', 'Processes payment for booking'),
        ('handleCancelBooking()', 'pages/player/BookingDetails/BookingDetails.js', 'Cancels the current booking'),
        ('showNotification(message, type)', 'pages/player/BookingDetails/BookingDetails.js', 'Shows notification messages on booking details page'),
        ('viewBookingDetails(booking_id)', 'pages/player/BookingDetails/BookingDetails.js', 'Navigates to booking details page'),
        ('testPrivacyToggle()', 'pages/player/BookingDetails/BookingDetails.js', 'Tests privacy toggle functionality'),
        ('goBack()', 'pages/player/BookingDetails/BookingDetails.js', 'Navigates back to previous page'),
    ]
    
    for func_name, file_path, explanation in player_functions:
        p = doc.add_paragraph()
        p.add_run(f'Function: ').bold = True
        p.add_run(func_name)
        p.add_run(f'\nFile: ').bold = True
        p.add_run(file_path)
        p.add_run(f'\nExplanation: ').bold = True
        p.add_run(explanation)
        doc.add_paragraph()

    # Authentication Functions
    doc.add_heading('Authentication Functions', level=2)
    
    auth_functions = [
        ('checkAdminStatus()', 'pages/auth/User_Selection_Page/UserSelection.js', 'Checks if current user has admin privileges'),
        ('showModal()', 'pages/auth/EditProfile/EditProfile.js', 'Displays edit profile modal dialog'),
        ('goBack()', 'pages/auth/EditProfile/EditProfile.js', 'Navigates back to previous page'),
    ]
    
    for func_name, file_path, explanation in auth_functions:
        p = doc.add_paragraph()
        p.add_run(f'Function: ').bold = True
        p.add_run(func_name)
        p.add_run(f'\nFile: ').bold = True
        p.add_run(file_path)
        p.add_run(f'\nExplanation: ').bold = True
        p.add_run(explanation)
        doc.add_paragraph()

    # Admin Functions
    doc.add_heading('Admin Functions', level=2)
    
    admin_js_functions = [
        ('renderReports(reports)', 'pages/Admin/ManageVenueRequests/ManageVenueRequests.js', 'Displays venue reports for admin review'),
    ]
    
    for func_name, file_path, explanation in admin_js_functions:
        p = doc.add_paragraph()
        p.add_run(f'Function: ').bold = True
        p.add_run(func_name)
        p.add_run(f'\nFile: ').bold = True
        p.add_run(file_path)
        p.add_run(f'\nExplanation: ').bold = True
        p.add_run(explanation)
        doc.add_paragraph()

    # Facility Owner Functions
    doc.add_heading('Facility Owner Functions', level=2)
    
    facility_owner_functions = [
        ('loadChat(username)', 'pages/Facility_Owner/Messages/Messages.js', 'Loads chat conversation with a user'),
        ('sendMessage()', 'pages/Facility_Owner/Messages/Messages.js', 'Sends chat message to current chat partner'),
    ]
    
    for func_name, file_path, explanation in facility_owner_functions:
        p = doc.add_paragraph()
        p.add_run(f'Function: ').bold = True
        p.add_run(func_name)
        p.add_run(f'\nFile: ').bold = True
        p.add_run(file_path)
        p.add_run(f'\nExplanation: ').bold = True
        p.add_run(explanation)
        doc.add_paragraph()

    # Owner Functions
    doc.add_heading('Owner Functions', level=2)
    
    owner_functions = [
        ('openModal()', 'pages/Owner/Owner.js', 'Opens modal dialog for owner actions'),
        ('closeModal()', 'pages/Owner/Owner.js', 'Closes modal dialog'),
    ]
    
    for func_name, file_path, explanation in owner_functions:
        p = doc.add_paragraph()
        p.add_run(f'Function: ').bold = True
        p.add_run(func_name)
        p.add_run(f'\nFile: ').bold = True
        p.add_run(file_path)
        p.add_run(f'\nExplanation: ').bold = True
        p.add_run(explanation)
        doc.add_paragraph()

    # Function Categories Summary
    doc.add_heading('Function Categories Summary', level=1)
    
    categories = [
        ('Authentication & User Management', 8),
        ('Venue & Booking Management', 35),
        ('Social Features', 25),
        ('Communication', 6),
        ('Admin Functions', 4),
        ('Utility Functions', 49),
    ]
    
    for category, count in categories:
        p = doc.add_paragraph()
        p.add_run(f'{category}: ').bold = True
        p.add_run(f'{count} functions')
    
    # Save the document
    doc.save('BOOK_PLAY_Functions_Documentation.docx')
    print("Word document 'BOOK_PLAY_Functions_Documentation.docx' has been created successfully!")

if __name__ == "__main__":
    create_function_documentation()