import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import glob

def create_word_document():
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('BOOK-PLAY Project - Complete Code Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add project description
    doc.add_paragraph('This document contains all the source code files from the BOOK-PLAY project.')
    doc.add_paragraph('Generated automatically for documentation purposes.')
    doc.add_paragraph('')
    
    # Define file extensions to include
    code_extensions = [
        '*.php', '*.html', '*.css', '*.js', '*.sql', '*.md', '*.txt',
        '*.py', '*.json', '*.xml', '*.yml', '*.yaml', '*.ini', '*.conf'
    ]
    
    # Get all files with code extensions
    all_files = []
    for ext in code_extensions:
        all_files.extend(glob.glob(f'**/{ext}', recursive=True))
    
    # Sort files for better organization
    all_files.sort()
    
    # Exclude certain directories
    excluded_dirs = ['.git', '__pycache__', 'node_modules', '.vscode']
    filtered_files = []
    
    for file_path in all_files:
        if not any(excluded in file_path for excluded in excluded_dirs):
            filtered_files.append(file_path)
    
    # Group files by directory
    files_by_dir = {}
    for file_path in filtered_files:
        dir_name = os.path.dirname(file_path) if os.path.dirname(file_path) else 'Root'
        if dir_name not in files_by_dir:
            files_by_dir[dir_name] = []
        files_by_dir[dir_name].append(file_path)
    
    # Add table of contents
    doc.add_heading('Table of Contents', level=1)
    toc_para = doc.add_paragraph()
    
    # Process each directory
    for dir_name in sorted(files_by_dir.keys()):
        # Add directory heading
        doc.add_heading(f'Directory: {dir_name}', level=1)
        
        # Add to table of contents
        toc_para.add_run(f'{dir_name}\n')
        
        # Process each file in the directory
        for file_path in sorted(files_by_dir[dir_name]):
            try:
                # Add file heading
                file_heading = doc.add_heading(f'File: {file_path}', level=2)
                
                # Add to table of contents
                toc_para.add_run(f'  - {file_path}\n')
                
                # Read and add file content
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                
                # Add file info
                file_info = doc.add_paragraph()
                file_info.add_run(f'File Size: {len(content)} characters\n')
                file_info.add_run(f'Lines: {len(content.splitlines())}\n')
                file_info.add_run(f'Extension: {os.path.splitext(file_path)[1]}\n')
                
                # Add content with syntax highlighting (basic)
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(content)
                code_run.font.name = 'Courier New'
                code_run.font.size = doc.styles['Normal'].font.size
                
                # Add separator
                doc.add_paragraph('=' * 80)
                doc.add_paragraph('')
                
            except Exception as e:
                # If there's an error reading the file, add a note
                error_para = doc.add_paragraph()
                error_para.add_run(f'Error reading file {file_path}: {str(e)}')
                doc.add_paragraph('')
    
    # Save the document
    output_file = 'BOOK-PLAY_Complete_Code_Documentation.docx'
    doc.save(output_file)
    print(f'Word document created successfully: {output_file}')
    print(f'Total files processed: {len(filtered_files)}')
    
    return output_file

if __name__ == "__main__":
    try:
        output_file = create_word_document()
        print(f"\nDocument saved as: {output_file}")
        print("You can now open this Word document to view all your project code.")
    except Exception as e:
        print(f"Error creating document: {e}") 