import os
from docx import Document
from docx.shared import Pt

# ✅ Folders you want to include all their files
target_folders = [
    r'C:\wamp64\www\BOOK-PLAY/pages/player',
    r'C:\wamp64\www\BOOK-PLAY/pages/Admin',
    r'C:\wamp64\www\BOOK-PLAY/pages/Facility_Owner',
    r'C:\wamp64\www\BOOK-PLAY/pages/auth',
    r'C:\wamp64\www\BOOK-PLAY/pages/Owner',
    r'C:\wamp64\www\BOOK-PLAY/assets/css',
    r'C:\wamp64\www\BOOK-PLAY/assets/js',
    r'C:\wamp64\www/BOOK-PLAY/components'
]

# ✅ Individual files (outside folders)
target_files = [
    r'C:\wamp64\www\BOOK-PLAY/mail/MailLink.php',
    r'C:\wamp64\www\BOOK-PLAY/admin_actions.php',
    r'C:\wamp64\www\BOOK-PLAY/index.php',
    r'C:\wamp64\www\BOOK-PLAY/logout.php'
]

# Desired file types
extensions = ['.php', '.js', '.html', '.css']

# Output Word file name
output_doc = 'Selected_Project_Code.docx'

# Create new Word document
doc = Document()
doc.add_heading('📘 Project Code Documentation (Selected Files & Folders)', 0)

# ✅ First: Iterate through specified folders
for folder in target_folders:
    for root, dirs, files in os.walk(folder):
        for file in files:
            if any(file.endswith(ext) for ext in extensions):
                filepath = os.path.join(root, file)
                try:
                    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                        code = f.read()

                    relative_path = os.path.relpath(filepath, folder)
                    doc.add_heading(relative_path, level=2)

                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(code)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8)

                    doc.add_paragraph('—' * 40)

                except Exception as e:
                    print(f'❌ Error reading {filepath}: {e}')

# ✅ Second: Read specified individual files
for filepath in target_files:
    if os.path.exists(filepath) and any(filepath.endswith(ext) for ext in extensions):
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                code = f.read()

            filename = os.path.basename(filepath)
            doc.add_heading(filename, level=2)

            paragraph = doc.add_paragraph()
            run = paragraph.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)

            doc.add_paragraph('—' * 40)

        except Exception as e:
            print(f'❌ Error reading {filepath}: {e}')
    else:
        print(f'⚠️ File not found or skipped (invalid extension): {filepath}')

# ✅ Save final Word file
doc.save(output_doc)
print(f"✅ Word file '{output_doc}' created successfully.")
